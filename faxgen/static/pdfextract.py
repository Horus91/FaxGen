from pypdf import PdfReader
from pathlib import Path
from datetime import datetime,timedelta
from docx import Document

OFPS_DIR= Path(__file__).resolve().parent.parent.joinpath('media').joinpath('ofps')
FXTEMP= Path(__file__).resolve().parent.parent.joinpath('static').joinpath('fxtemp.docx')


class Fax_elts():
  
    def __init__(self, fileurl):
        self.file=fileurl
        reader=PdfReader(fileurl)
        self.pages=reader.pages
        self.garde=reader.pages[0].extract_text()
        self.first=reader.pages[1].extract_text()

        # -----------Time Properties----------------
        self.datetime_dep=datetime.strptime(self.first[198:205]+self.first[676:680],'%d%b%y%H%M')
        self.trip_time= timedelta(hours=int(self.first[1267:1269]),minutes=int(self.first[1270:1272]))
        self.datetime_arr=self.datetime_dep+self.trip_time
        # ------------------------------------------


    def general_infos(self):
        reg_marks=self.first[655:660]
        callsign=self.garde[14:21]
        origin=self.first[759:786]
        destination=self.first[829:856]
        alt1=self.first[899:926]
        alt2=self.first[969:996]
        return reg_marks,callsign,origin,destination,alt1,alt2
    
    def firs_extraction(self):
        for page in self.pages:
            e_page=page.extract_text()
            if '(FPL-' in e_page:
                firs=e_page[e_page.index('EET/')+4: e_page.index('SEL/')-1].split(' ')
                firs_lst=[]
                firs_timing=[]
                for fir in firs:
                    firs_lst.append(fir[:4])
                    firs_timing.append(fir[4:8])
                global last_log_page
                last_log_page=page.page_number

                return firs_lst,firs_timing
        
    def route_extraction(self):
        for page in self.pages:
            if 'ATC ROUTING:' in page.extract_text():
                atc_id=page.extract_text().find('ATC ROUTING:')
                rvsm_id=page.extract_text().find('RVSM')

                return page.extract_text()[atc_id+18:rvsm_id-75].replace("\n"," ")
            
    def __str__(self):
        return f"{self.general_infos()[0]}, {self.general_infos()[2]}-{self.general_infos()[3]}, EET:{self.trip_time}"


class Fax_docx():
    def __init__(self,fir_countries,firs_points,ac_type,captain,Fax_elts):
    # fir_countries,firs_points,captain,
        self.fax_template=Document(FXTEMP)
        self.fir_countries= fir_countries
        self.fir_points=firs_points
        self.ac_type=ac_type
        self.captain=captain
        self.Fax_elts=Fax_elts
        self.general_infos=self.Fax_elts.general_infos()
        


    def generate_fax(self):
        # Cells positions {'5':pays survoles, '17':nbre type d'avion, '34': callsign, '34':registration, '44':Captain
        # , '120':Depdate,'125':etd,'122':Dep_airport,,'127':Arr_airport, '129':eta, '131':route,
        #  }
        table=self.fax_template.tables[0]
        table1=self.fax_template.tables[0]._cells[159].tables[0]
        
        
        # General Infos Respectively (fir_countries;type;captain;reg;callsign;origin;destination)
        table._cells[5].text = self.fir_countries
        table._cells[17].text = self.ac_type
        table._cells[44].text = self.captain
        table._cells[34].text = self.general_infos[0] 
        table._cells[24].text = self.general_infos[1]
        table._cells[122].text = self.general_infos[2]
        table._cells[127].text = self.general_infos[3]

        # Date Times 
        table._cells[120].text = self.Fax_elts.datetime_dep.strftime("%d%b%Y")
        table._cells[125].text = self.Fax_elts.datetime_dep.strftime("%H:%M GMT")
        table._cells[129].text = self.Fax_elts.datetime_arr.strftime("%H:%M GMT\n%d%b%Y")

        # Route:
        table._cells[131].text = self.Fax_elts.route_extraction()
        
        # FIRS:
        row_firs=table.rows[-1]
        firs_extracted=self.Fax_elts.firs_extraction()[0]
        firs_etos=self.Fax_elts.firs_extraction()[1]
        firs_pts=self.fir_points.split(',')
        while(len(firs_extracted)>len(firs_pts)):
            firs_pts.append(' ')
        a=0
        for elt in range(len(firs_extracted)):
            a+=1
            table.add_row()._tr.addnext(row_firs._tr)
            table.rows[14+a].cells[0].merge(table.rows[14+a].cells[2])
            table.rows[14+a].cells[3].merge(table.rows[14+a].cells[7])
            table.rows[14+a].cells[8].merge(table.rows[14+a].cells[9])
            table.rows[14+a].cells[0].text=firs_extracted[elt]
            table.rows[14+a].cells[3].text=firs_pts[elt]
            table.rows[14+a].cells[8].text=f"H+ {firs_etos[elt]}"
        
        # ALTERNATES:
            table1._cells[1].text=self.Fax_elts.general_infos()[4]
            table1._cells[3].text=self.Fax_elts.general_infos()[5]
        
        return self.fax_template.save(OFPS_DIR.joinpath(f'{self.general_infos[0]}-{self.captain}-{self.general_infos[2][:4]}-{self.general_infos[3][:4]}.docx'))
        

        
        
# b=Fax_docx(Fax_elts(OFPS_DIR.joinpath('FlightPlan_K5461.pdf')))
# b.generate_fax()
    


# t=['a','b','c','d']
# new_fax=Document(FXTEMP)
# table= new_fax.tables[0]._cells[159].tables[0]
# print(table._cells[3].text)
# print(new_fax.tables[0]._cells[44].text)